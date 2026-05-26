"""
output_generator.py
Generates Excel output strictly matching One_TC_Updated.xlsx template format.

Exact column layout (matches template):
  A(1)   Requirement_ID
  B(2)   TC_ID
  C(3)   Scenario No          -- format: SC_001, SC_002 ...
  D(4)   Test Objective
  E(5)   Test Details Description
  F(6)   Test Precondition     -- Req 5: consolidates E + input-related Test Steps from H col
  G(7)   Inputs               -- merged header over sub-signal columns H, I, ...
  H(8)+  [input signal sub-columns, dynamic]
  J(10)  Test Steps           -- standalone column (after input sub-cols)
  K(11)  Expected Outputs     -- merged header over output signal sub-columns
  L(12)+ [output signal sub-columns, dynamic]
  M(13)  Depands On           -- TC_ID + Scenario No concatenated (Req 10)
  N(14)  Test_Env
  O(15)  Test_Type
  P(16)  Scenario_Type
  Q(17)  Remarks/Additional information  -- bullet format, no test-basis (Req 8)
  R(18)  Module               -- alpha-only (Req 7)

  NOTE: Column positions G/J/K/M etc. shift right if there are more input/output signals.
  The template example has 3 input signals and 2 output signals, giving:
    G(7)=Inputs header, H(8)=sig1, I(9)=sig2, [J(10)=sig3 if 3 inputs]
  Since signals are dynamic, we compute offsets at runtime.
"""

import io
import re
from datetime import datetime
from typing import List, Dict, Tuple

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import TestCase
from config import ENGINE


# ─── STYLING — Uniform colour scheme (Requirement 6) ─────────────────────────
# All header cells use the same blue fill; no per-column different colours.
HEADER_FILL   = PatternFill("solid", fgColor="4472C4")   # uniform blue
HEADER_FONT   = Font(bold=True, color="FFFFFF", size=10, name="Calibri")
HEADER_ALIGN  = Alignment(horizontal="center", vertical="center", wrap_text=True)

SUBHDR_FILL   = PatternFill("solid", fgColor="4472C4")   # same blue for sub-headers (Req 6)
SUBHDR_FONT   = Font(bold=True, color="FFFFFF", size=9,  name="Calibri")
SUBHDR_ALIGN  = Alignment(horizontal="center", vertical="center", wrap_text=True)

BODY_FONT     = Font(size=9, name="Calibri")
BODY_ALIGN    = Alignment(vertical="top", wrap_text=True)
CENTER_ALIGN  = Alignment(horizontal="center", vertical="top", wrap_text=True)

THIN_SIDE     = Side(style="thin", color="CCCCCC")
THIN_BORDER   = Border(left=THIN_SIDE, right=THIN_SIDE, top=THIN_SIDE, bottom=THIN_SIDE)

ALT_FILL      = PatternFill("solid", fgColor="EEF2F9")   # alternating row shading


# ─── SIGNAL EXTRACTION ────────────────────────────────────────────────────────

_KV = re.compile(r'^(.+?):\s*(.+)$')


def _parse_signal_value(entry: str) -> Tuple[str, str]:
    m = _KV.match(entry.strip())
    return (m.group(1).strip(), m.group(2).strip()) if m else (entry.strip(), entry.strip())


def extract_signal_columns(test_cases: List[TestCase]) -> Tuple[List[str], List[str]]:
    """Return ordered unique input signal names and output signal names across all TCs."""
    in_sigs, out_sigs = [], []
    seen_in, seen_out = set(), set()

    for tc in test_cases:
        for entry in tc.inputs:
            name, _ = _parse_signal_value(entry)
            if name and name not in seen_in:
                seen_in.add(name); in_sigs.append(name)

        for m in re.finditer(r'([\w\s]{3,50}?)\s*[=:]\s*(\w+)', tc.expected_outcome):
            name = m.group(1).strip()
            if any(s in name.lower() for s in [
                "system successfully", "response is", "data is", "result is",
                "all sub", "no data", "logic module", "specification",
                "test case", "is correct", "the logic", "and sets"
            ]):
                continue
            if len(name.split()) <= 6 and name not in seen_out:
                seen_out.add(name); out_sigs.append(name)

    return in_sigs, out_sigs


def _get_signal_value(tc: TestCase, signal_name: str, kind: str) -> str:
    if kind == "input":
        for entry in tc.inputs:
            name, value = _parse_signal_value(entry)
            if name.lower() == signal_name.lower():
                return value
        return ""
    for m in re.finditer(r'([\w\s]{3,50}?)\s*[=:]\s*(\w+)', tc.expected_outcome):
        if m.group(1).strip().lower() == signal_name.lower():
            return m.group(2).strip()
    return ""


# ─── FIELD HELPERS ────────────────────────────────────────────────────────────

def _list_to_str(value) -> str:
    if isinstance(value, list):
        return "\n".join(str(v) for v in value if v)
    return str(value) if value else ""


def _cell_value(tc: TestCase, field: str) -> str:
    return _list_to_str(getattr(tc, field, ""))


def _module_alpha_only(module: str) -> str:
    """Requirement 7: keep only alphabetical characters and spaces."""
    cleaned = re.sub(r'[^A-Za-z\s]', '', module).strip()
    return re.sub(r'\s+', ' ', cleaned) or "General"


def _sc_label(sc_no: int) -> str:
    """Format scenario number as SC_001, SC_002, etc. to match template."""
    return f"SC_{sc_no:03d}"


# ─── REQUIREMENT 5: Column F content ─────────────────────────────────────────
# Col F = Test Precondition, but per Req 5 it must consolidate:
#   • Test Objective (from col D/E)
#   • Test Steps that are related to the identified input parameter names (from H, I... cols)

def _col_f_precondition(tc: TestCase, input_signals: List[str]) -> str:
    """
    Requirement 5: Column F consolidates Test Objective + input-related Test Steps.
    Also satisfies Requirement 9: includes pre-set input values and output-influence notes.
    """
    parts = []

    # --- Test Objective (from col D) ---
    parts.append(f"Test Objective:\n{tc.objective}")

    # --- Pre-set input conditions from requirement (Req 9) ---
    preset_vals = re.findall(
        r'(?:is|=|equals?|set\s+to)\s+(True|False|Active|Inactive|\d+(?:\.\d+)?)',
        tc.objective + " " + " ".join(tc.preconditions),
        re.IGNORECASE
    )
    if preset_vals:
        parts.append(
            "Pre-set input values (from requirement): " + ", ".join(preset_vals) + "\n"
            "Output is generated when ALL specified input conditions are met. "
            "Any change in the above inputs will directly influence the output."
        )

    # --- Input-related Test Steps (from col H / input sub-columns) ---
    # Collect input names to match against steps
    input_names = {s.lower() for s in input_signals}
    for entry in tc.inputs:
        name, _ = _parse_signal_value(entry)
        if name:
            input_names.add(name.lower())

    matched_steps = []
    steps = tc.test_steps if isinstance(tc.test_steps, list) else [tc.test_steps]
    for step in steps:
        sl = step.lower()
        if any(n in sl for n in input_names) or any(
            kw in sl for kw in ['input', 'prepare', 'set ', 'boundary', 'valid', 'value']
        ):
            matched_steps.append(step)

    if matched_steps:
        parts.append("Test Steps (input-related):\n" + "\n".join(matched_steps))

    # --- Standard preconditions ---
    if tc.preconditions:
        std_pre = _list_to_str(tc.preconditions)
        parts.append("Preconditions:\n" + std_pre)

    return "\n\n".join(parts)


# ─── REQUIREMENT 8: Remarks bullet formatting ─────────────────────────────────

def _remarks_bullets(tc: TestCase) -> str:
    """
    Requirement 8:
    - Remove test-basis-related info
    - Include type of testing per scenario
    - Describe what is tested in each SC (e.g. INPUT_1 maximum value is tested)
    - Bullet-point format
    """
    bullets = []

    # Type of testing for this scenario
    bullets.append(f"• Testing Type: {tc.testing_type.capitalize()} | Scenario Type: {tc.scenario_type.capitalize()}")

    # What is being tested (Req 8 — describe each SC)
    sc_what = {
        "normal":     "All input values set to normal/valid values; correct system output is verified.",
        "boundary":   "Input boundary values tested: minimum, maximum, min-1, max+1 for each parameter.",
        "edge":       "Edge case conditions tested (state transitions, simultaneous changes, unusual-but-valid states).",
        "robustness": "Invalid/out-of-range input values tested; system must respond safely without crash.",
    }
    bullets.append(f"• What is tested: {sc_what.get(tc.scenario_type, 'Functional system behaviour verified.')}")

    # Per-input description (e.g. "INPUT_1 maximum value is tested")
    for entry in tc.inputs:
        name, value = _parse_signal_value(entry)
        if name and value and name.lower() not in ("test environment", "all prerequisite", "sub-requirements"):
            if tc.scenario_type == "boundary":
                if "max" in value.lower() or "maximum" in value.lower():
                    bullets.append(f"• {name}: maximum value is tested")
                elif "min" in value.lower() or "minimum" in value.lower():
                    bullets.append(f"• {name}: minimum value is tested")
                elif "-1" in value or "below" in value.lower():
                    bullets.append(f"• {name}: below-minimum value is tested (invalid range)")
                elif "+1" in value or "above" in value.lower():
                    bullets.append(f"• {name}: above-maximum value is tested (invalid range)")
                else:
                    bullets.append(f"• {name}: boundary value '{value}' is tested")
            elif tc.scenario_type == "edge":
                bullets.append(f"• {name}: edge-case value '{value}' is tested (state-transition condition)")
            elif tc.scenario_type == "robustness":
                bullets.append(f"• {name}: invalid/out-of-range value '{value}' is tested")

    # Input source note (Req 4)
    inputs_raw = " ".join(tc.inputs).lower()
    if any(kw in inputs_raw for kw in ["icd", "derived", "interface"]):
        bullets.append("• Input source: Values derived from ICD document (not explicitly defined in SRS).")
    else:
        bullets.append("• Input source: Input values explicitly defined in SRS specification.")

    # Sub-requirements / cross-refs from raw remarks (strip test-basis lines)
    if tc.remarks:
        raw_parts = re.split(r'\s*[\|\n•]+\s*', tc.remarks)
        for part in raw_parts:
            part = part.strip()
            if not part:
                continue
            # Remove test-basis lines (Req 8)
            if re.search(
                r'test\s+basis|input\s+values\s+derived\s+from\s+srs|srs\s+requirement\s+\w',
                part, re.IGNORECASE
            ):
                continue
            # Include enum definitions, sub-req refs, notes
            if re.search(r'enum|sub.req|note|reference|derived from icd|document context', part, re.IGNORECASE):
                bullets.append(f"• {part}")

    return "\n".join(bullets)


# ─── REQUIREMENT 10: Depends On ───────────────────────────────────────────────

def _depends_on(raw_dep: str, tc_id: str, sc_no: int) -> str:
    """
    Depands On column.
    Format: TC_UT_001_SC-001  (hyphen between SC and number)

    The generator writes:
      - "None"            for SC_001 (baseline)
      - "TC_UT_001_SC-001" for SC_002+ (always references baseline with hyphen)

    This function passes the value through unchanged if already formatted,
    or applies a fallback for legacy/MCP data.
    """
    if not raw_dep or raw_dep.strip().lower() == "none":
        return "None"
    raw = raw_dep.strip()
    # Already formatted (TC_ID_SC-001 hyphen format or TC_ID_SC_001 underscore)
    if "_SC-" in raw or "_SC_" in raw.upper():
        return raw
    # Fallback: bare TC_ID — append SC-001 (baseline reference, hyphen format)
    return f"{raw}_SC-001"


# ─── HEADER WRITER ────────────────────────────────────────────────────────────

def _write_headers(ws, input_signals: List[str], output_signals: List[str]) -> Dict[str, int]:
    """
    Writes rows 1 and 2 exactly matching One_TC_Updated.xlsx template.
    Returns a dict of column-name -> column-index for use when writing data.

    Template exact layout:
      Col 1: Requirement_ID  (rows 1-2 merged)
      Col 2: TC_ID           (rows 1-2 merged)
      Col 3: Scenario No     (rows 1-2 merged)
      Col 4: Test Objective  (rows 1-2 merged)
      Col 5: Test Details Description  (rows 1-2 merged)
      Col 6: Test Precondition         (rows 1-2 merged)
      Col 7: Inputs          (row 1 merged across input signal sub-cols)
        Col 7+0: signal_1 sub-header (row 2)
        Col 7+1: signal_2 sub-header (row 2)
        ...
      Col 7+n_inputs: Test Steps       (rows 1-2 merged)
      Col 7+n_inputs+1: Expected Outputs (row 1 merged across output sub-cols)
        output sub-headers (row 2)
      Col 7+n_inputs+1+n_outputs: Depands On   (rows 1-2 merged)  [sic]
      Col ...: Test_Env       (rows 1-2 merged)
      Col ...: Test_Type      (rows 1-2 merged)
      Col ...: Scenario_Type  (rows 1-2 merged)
      Col ...: Remarks/Additional information  (rows 1-2 merged)
      Col ...: Module          (rows 1-2 merged)
    """
    n_in  = len(input_signals)
    n_out = len(output_signals)

    # Fixed prefix columns A-F
    prefix = [
        ("Requirement_ID",          21),
        ("TC_ID",                    9),
        ("Scenario No",             12),
        ("Test Objective",          20),
        ("Test Details Description",22),
        ("Test Precondition",       45),
    ]

    col = 1
    col_map: Dict[str, int] = {}

    # Write prefix headers (each spans rows 1-2)
    for hdr, width in prefix:
        c = ws.cell(row=1, column=col, value=hdr)
        c.font = HEADER_FONT; c.fill = HEADER_FILL
        c.alignment = HEADER_ALIGN; c.border = THIN_BORDER
        ws.column_dimensions[get_column_letter(col)].width = width
        ws.merge_cells(start_row=1, start_column=col, end_row=2, end_column=col)
        col_map[hdr] = col
        col += 1

    # "Inputs" group header at col G
    inputs_start = col
    col_map["Inputs_start"] = inputs_start
    c = ws.cell(row=1, column=col, value="Inputs")
    c.font = HEADER_FONT; c.fill = HEADER_FILL
    c.alignment = HEADER_ALIGN; c.border = THIN_BORDER
    if n_in > 1:
        ws.merge_cells(start_row=1, start_column=col, end_row=1, end_column=col + n_in - 1)
    elif n_in == 0:
        ws.merge_cells(start_row=1, start_column=col, end_row=2, end_column=col)

    # Input signal sub-headers in row 2
    for i, sig in enumerate(input_signals):
        c2 = ws.cell(row=2, column=col + i, value=sig)
        c2.font = SUBHDR_FONT; c2.fill = SUBHDR_FILL   # same blue (Req 6)
        c2.alignment = SUBHDR_ALIGN; c2.border = THIN_BORDER
        ws.column_dimensions[get_column_letter(col + i)].width = max(18, len(sig) + 4)
        col_map[f"input_sig_{i}"] = col + i
    col += max(n_in, 1)  # advance at least 1 column

    # "Test Steps" standalone column
    col_map["Test Steps"] = col
    c = ws.cell(row=1, column=col, value="Test Steps")
    c.font = HEADER_FONT; c.fill = HEADER_FILL
    c.alignment = HEADER_ALIGN; c.border = THIN_BORDER
    ws.column_dimensions[get_column_letter(col)].width = 30
    ws.merge_cells(start_row=1, start_column=col, end_row=2, end_column=col)
    col += 1

    # "Expected Outputs" group header
    outputs_start = col
    col_map["Outputs_start"] = outputs_start
    c = ws.cell(row=1, column=col, value="Expected Outputs")
    c.font = HEADER_FONT; c.fill = HEADER_FILL
    c.alignment = HEADER_ALIGN; c.border = THIN_BORDER
    if n_out > 1:
        ws.merge_cells(start_row=1, start_column=col, end_row=1, end_column=col + n_out - 1)
    elif n_out == 0:
        ws.merge_cells(start_row=1, start_column=col, end_row=2, end_column=col)

    for i, sig in enumerate(output_signals):
        c2 = ws.cell(row=2, column=col + i, value=sig)
        c2.font = SUBHDR_FONT; c2.fill = SUBHDR_FILL   # same blue (Req 6)
        c2.alignment = SUBHDR_ALIGN; c2.border = THIN_BORDER
        ws.column_dimensions[get_column_letter(col + i)].width = max(22, len(sig) + 4)
        col_map[f"output_sig_{i}"] = col + i
    col += max(n_out, 1)

    # Suffix columns — all same blue header (Req 6)
    suffix = [
        ("Depands On",                      12),   # sic — typo preserved from template
        ("Test_Env",                        12),
        ("Test_Type",                       16),
        ("Scenario_Type",                   14),
        ("Remarks/Additional information",  32),
        ("Module",                           9),
    ]
    for hdr, width in suffix:
        c = ws.cell(row=1, column=col, value=hdr)
        c.font = HEADER_FONT; c.fill = HEADER_FILL
        c.alignment = HEADER_ALIGN; c.border = THIN_BORDER
        ws.column_dimensions[get_column_letter(col)].width = width
        ws.merge_cells(start_row=1, start_column=col, end_row=2, end_column=col)
        col_map[hdr] = col
        col += 1

    ws.row_dimensions[1].height = 28
    ws.row_dimensions[2].height = 22
    ws.freeze_panes = "A3"
    return col_map


# ─── EXCEL EXPORT ─────────────────────────────────────────────────────────────

def generate_excel(test_cases: List[TestCase], removed_count: int) -> bytes:
    """
    Generate Excel matching One_TC_Updated.xlsx template exactly.
    All requirements applied:
      Req 3:  TC_ID same for all scenarios of one req; SC resets per req
      Req 4:  Input source (SRS/ICD) recorded in Remarks
      Req 5:  Col F = Test Objective + input-related Test Steps
      Req 6:  Uniform blue header colour throughout
      Req 7:  Module = alpha-only
      Req 8:  Remarks = bullet format, no test-basis, SC description
      Req 9:  Precondition includes pre-set values + output-influence note
      Req 10: Depands On = TC_ID + SC_NNN
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "test_cases"

    input_signals, output_signals = extract_signal_columns(test_cases)
    col_map = _write_headers(ws, input_signals, output_signals)

    # TC_ID, SC are already correctly assigned by _resequence() in test_case_generator.
    # Use them directly — no re-grouping needed here.
    for row_idx, tc in enumerate(test_cases, start=3):
        is_alt = (row_idx % 2 == 0)
        req_id = tc.traceability_req_id
        tc_id  = tc.test_case_id
        sc_lbl = tc.scenario_id
        sc_no  = int(sc_lbl.replace("SC_", "")) if sc_lbl.startswith("SC_") else row_idx - 2

        def _put(col: int, value, align=BODY_ALIGN):
            cell = ws.cell(row=row_idx, column=col, value=value)
            cell.font   = BODY_FONT
            cell.alignment = align
            cell.border = THIN_BORDER
            if is_alt:
                cell.fill = ALT_FILL

        # Fixed prefix columns
        _put(col_map["Requirement_ID"],          req_id)
        _put(col_map["TC_ID"],                   tc_id)
        _put(col_map["Scenario No"],             sc_lbl)
        _put(col_map["Test Objective"],          tc.objective)
        _put(col_map["Test Details Description"],_list_to_str(tc.preconditions))  # Col E = detail context
        _put(col_map["Test Precondition"],       _col_f_precondition(tc, input_signals))  # Col F (Req 5)

        # Input signal sub-columns
        for i, sig in enumerate(input_signals):
            val = _get_signal_value(tc, sig, "input")
            _put(col_map["Inputs_start"] + i, val, CENTER_ALIGN)

        # Test Steps (standalone column J)
        steps_str = _list_to_str(tc.test_steps)
        _put(col_map["Test Steps"], steps_str)

        # Output signal sub-columns
        for i, sig in enumerate(output_signals):
            val = _get_signal_value(tc, sig, "output")
            _put(col_map["Outputs_start"] + i, val, CENTER_ALIGN)

        # Depands On (Req 10: TC_ID + Scenario No)
        _put(col_map["Depands On"], _depends_on(tc.dependent_test_cases, tc_id, sc_no))

        # Test_Env, Test_Type, Scenario_Type
        _put(col_map["Test_Env"],      tc.test_environment)
        _put(col_map["Test_Type"],     tc.testing_type)
        _put(col_map["Scenario_Type"], tc.scenario_type)

        # Remarks (Req 8: bullet format, no test-basis)
        _put(col_map["Remarks/Additional information"], _remarks_bullets(tc))

        # Module (Req 7: alpha-only)
        _put(col_map["Module"], _module_alpha_only(tc.module))

    # ── Summary sheet ─────────────────────────────────────────────────────────
    ws2 = wb.create_sheet(title="Summary")
    ws2.column_dimensions["A"].width = 35
    ws2.column_dimensions["B"].width = 25

    sum_hdr_font = Font(bold=True, color="FFFFFF", size=11, name="Calibri")
    sum_hdr_fill = PatternFill("solid", fgColor="2F4F8F")
    lbl_font     = Font(bold=True, size=10, name="Calibri")
    val_font     = Font(size=10, name="Calibri")

    def _sh_title(r, text):
        c = ws2.cell(row=r, column=1, value=text)
        c.font = sum_hdr_font; c.fill = sum_hdr_fill
        ws2.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
        c.alignment = Alignment(horizontal="center")

    def _sh_row(r, label, value):
        ws2.cell(row=r, column=1, value=label).font = lbl_font
        ws2.cell(row=r, column=2, value=value).font = val_font

    from collections import Counter
    r = 1
    _sh_title(r, "Test Case Generation Summary"); r += 1
    _sh_row(r, "Total Test Cases", len(test_cases)); r += 1
    _sh_row(r, "Duplicates Removed", removed_count); r += 1
    _sh_row(r, "Generated On", datetime.now().strftime("%Y-%m-%d %H:%M:%S")); r += 2

    _sh_title(r, "By Module"); r += 1
    for mod, cnt in sorted(Counter(_module_alpha_only(tc.module) for tc in test_cases).items()):
        _sh_row(r, mod, cnt); r += 1
    r += 1

    _sh_title(r, "By Scenario Type"); r += 1
    for st, cnt in sorted(Counter(tc.scenario_type for tc in test_cases).items()):
        _sh_row(r, st.capitalize(), cnt); r += 1
    r += 1

    _sh_title(r, "By Testing Type"); r += 1
    for tt, cnt in sorted(Counter(tc.testing_type for tc in test_cases).items()):
        _sh_row(r, tt.capitalize(), cnt); r += 1

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─── WORD EXPORT ──────────────────────────────────────────────────────────────

def generate_docx(test_cases: List[TestCase], removed_count: int) -> bytes:
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.8)
        section.left_margin = section.right_margin = Inches(0.9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Test Case Report")
    run.font.size = Pt(20); run.font.bold = True
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        f"Total: {len(test_cases)} test cases  |  Duplicates removed: {removed_count}"
    ).font.size = Pt(9)
    doc.add_paragraph()

    from collections import defaultdict
    by_module = defaultdict(list)
    for tc in test_cases:
        by_module[_module_alpha_only(tc.module)].append(tc)

    input_signals, output_signals = extract_signal_columns(test_cases)

    for module in sorted(by_module.keys()):
        h = doc.add_paragraph(f"Module: {module}")
        h.style = "Heading 1"

        for tc in by_module[module]:
            req_id = tc.traceability_req_id
            tc_id  = tc.test_case_id
            sc_lbl = tc.scenario_id
            sc_no  = int(sc_lbl.replace("SC_", "")) if sc_lbl.startswith("SC_") else 1

            sub = doc.add_paragraph(f"{tc_id} | {sc_lbl} | {tc.scenario_type.capitalize()}")
            sub.style = "Heading 2"

            rows = [
                ("Requirement_ID",          req_id),
                ("TC_ID",                   tc_id),
                ("Scenario No",             sc_lbl),
                ("Test Objective",          tc.objective),
                ("Test Details Description",_list_to_str(tc.preconditions)),
                ("Test Precondition",       _col_f_precondition(tc, input_signals)),
                ("Inputs",                  _list_to_str(tc.inputs)),
                ("Test Steps",              _list_to_str(tc.test_steps)),
                ("Expected Outputs",        tc.expected_outcome),
                ("Depands On",              _depends_on(tc.dependent_test_cases, tc_id, sc_no)),
                ("Test_Env",                tc.test_environment),
                ("Test_Type",               tc.testing_type),
                ("Scenario_Type",           tc.scenario_type),
                ("Remarks",                 _remarks_bullets(tc)),
                ("Module",                  _module_alpha_only(tc.module)),
            ]

            table = doc.add_table(rows=len(rows), cols=2)
            table.style = "Table Grid"
            for ri, (label, val) in enumerate(rows):
                row = table.rows[ri]
                lc = row.cells[0]; lc.width = Inches(2.0)
                lr = lc.paragraphs[0].add_run(label)
                lr.font.bold = True; lr.font.size = Pt(9)
                vr = row.cells[1].paragraphs[0].add_run(str(val))
                vr.font.size = Pt(9)

            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()