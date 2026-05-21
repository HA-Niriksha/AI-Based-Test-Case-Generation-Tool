import io
import fitz  # PyMuPDF
from docx import Document
import openpyxl


def parse_pdf(file_bytes: bytes) -> str:
    parts = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                parts.append(f"[Page {page_num + 1}]\n{text}")
        doc.close()
    except Exception as e:
        raise RuntimeError(f"PDF parsing error: {e}")
    return "\n".join(parts)


def _paragraph_is_bold_heading(para) -> bool:
    """
    Returns True if an entire paragraph is bold-formatted (module heading).
    Checks both paragraph-level bold runs and character-level formatting.
    """
    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return False
    # All non-empty runs must be bold
    return all(r.bold for r in runs)


def parse_docx(file_bytes: bytes) -> str:
    """
    Parses a DOCX file, preserving:
    - Bold headings as [MODULE: <text>] markers for module detection
    - Normal style headings as ## markers
    - Tables with pipe-separated rows
    Requirement: per spec §5.3, bold-formatted headings define module boundaries.
    """
    parts = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                parts.append(f"\n## {text}")
                parts.append(f"[MODULE: {text}]")
            elif _paragraph_is_bold_heading(para):
                # Bold non-heading paragraph = module section boundary
                parts.append(f"\n## {text}")
                parts.append(f"[MODULE: {text}]")
            else:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
    except Exception as e:
        raise RuntimeError(f"DOCX parsing error: {e}")
    return "\n".join(parts)


def parse_xlsx(file_bytes: bytes) -> str:
    parts = []
    try:
        wb = openpyxl.load_workbook(
            io.BytesIO(file_bytes), read_only=True, data_only=True
        )
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"\n[Sheet: {sheet_name}]")
            headers = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if all(v is None for v in row):
                    continue
                if i == 0:
                    headers = [str(c).strip() if c is not None else f"Col{j}" for j, c in enumerate(row)]
                    continue
                for h, val in zip(headers, row):
                    if val is not None and str(val).strip():
                        parts.append(f"{h}: {val}")
        wb.close()
    except Exception as e:
        raise RuntimeError(f"Excel parsing error: {e}")
    return "\n".join(parts)


def parse_file(filename: str, file_bytes: bytes) -> str:
    if not filename or not file_bytes:
        raise ValueError("Filename and file content are required")
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return parse_docx(file_bytes)
    elif ext in ("xlsx", "xls"):
        return parse_xlsx(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type: .{ext}. Accepted: .pdf, .docx, .xlsx"
        )
